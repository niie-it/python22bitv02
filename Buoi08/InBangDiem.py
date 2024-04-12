'''
Sử dụng thư viện python-docx và dữ liệu sẵn có in ra bảng điểm
Format bảng điểm:
Họ tên:____________________
|STT    | Tên môn   | Điểm  |
|1      | Python    | 9     |
'''
data = [
    {
        "ho_ten": "Tèo Nguyễn",
        "mon_hoc": [
            { "ma_mon": "MH01", "ten_mon": "Python", "diem": 9.1},
            { "ma_mon": "MH02", "ten_mon": "Security", "diem": 8.8},
        ]
    }
]

from docx import Document
document = Document() # Tạo file mới

document.add_heading("Họ tên: " + data[0]["ho_ten"], level=0)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'STT'
hdr_cells[1].text = 'Tên môn'
hdr_cells[2].text = 'Điểm'
for mon in data[0]["mon_hoc"]:
    row_cells = table.add_row().cells
    row_cells[0].text = mon["ma_mon"]
    row_cells[1].text = mon["ten_mon"]
    row_cells[2].text = str(mon["diem"])
    
document.save("BangDiem.docx") # Lưu